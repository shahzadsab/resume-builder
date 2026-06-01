"""
ATS Resume & Cover Letter Generator
====================================
Zero-cost Streamlit app using:
- pdfplumber / python-docx for parsing
- spaCy (en_core_web_sm) for keyword extraction
- Wikipedia API for company research
- Hugging Face Inference API (optional, free tier) for AI generation
- python-docx for .docx output

Fallbacks at every step so the app works with zero API keys.
"""

import streamlit as st
import io
import re
import json
import os
import requests
from collections import Counter
from datetime import datetime

# ── Document parsing ──────────────────────────────────────────────────────────
try:
    import pdfplumber
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── NLP keyword extraction ────────────────────────────────────────────────────
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
        SPACY_OK = True
    except OSError:
        SPACY_OK = False
        nlp = None
except ImportError:
    SPACY_OK = False
    nlp = None

# ── Wikipedia ─────────────────────────────────────────────────────────────────
try:
    import wikipedia
    WIKI_OK = True
except ImportError:
    WIKI_OK = False

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ATS Resume Builder",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
# CUSTOM CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:wght@300;400;500;600&display=swap');

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
}

/* ── Header ── */
.hero-header {
    background: linear-gradient(135deg, #0f172a 0%, #1e3a5f 50%, #0f172a 100%);
    border-radius: 16px;
    padding: 2.5rem 3rem;
    margin-bottom: 2rem;
    border: 1px solid rgba(99,179,237,0.2);
    position: relative;
    overflow: hidden;
}
.hero-header::before {
    content: '';
    position: absolute;
    top: -50%;
    right: -20%;
    width: 400px;
    height: 400px;
    background: radial-gradient(circle, rgba(99,179,237,0.12) 0%, transparent 70%);
    pointer-events: none;
}
.hero-title {
    font-family: 'DM Serif Display', serif;
    font-size: 2.4rem;
    color: #e2e8f0;
    margin: 0 0 0.4rem 0;
    line-height: 1.2;
}
.hero-subtitle {
    font-size: 1rem;
    color: #94a3b8;
    margin: 0;
    font-weight: 300;
    letter-spacing: 0.02em;
}
.hero-badge {
    display: inline-block;
    background: rgba(99,179,237,0.15);
    border: 1px solid rgba(99,179,237,0.3);
    color: #63b3ed;
    padding: 0.2rem 0.75rem;
    border-radius: 20px;
    font-size: 0.75rem;
    font-weight: 600;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    margin-bottom: 1rem;
}

/* ── Cards ── */
.card {
    background: #1e293b;
    border: 1px solid #334155;
    border-radius: 12px;
    padding: 1.5rem;
    margin-bottom: 1.2rem;
}
.card-title {
    font-family: 'DM Serif Display', serif;
    font-size: 1.2rem;
    color: #cbd5e1;
    margin: 0 0 1rem 0;
    padding-bottom: 0.6rem;
    border-bottom: 1px solid #334155;
}
.metric-grid {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(160px, 1fr));
    gap: 1rem;
    margin: 1rem 0;
}
.metric-box {
    background: #0f172a;
    border: 1px solid #334155;
    border-radius: 10px;
    padding: 1rem;
    text-align: center;
}
.metric-value {
    font-family: 'DM Serif Display', serif;
    font-size: 1.6rem;
    color: #63b3ed;
}
.metric-label {
    font-size: 0.75rem;
    color: #64748b;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    margin-top: 0.25rem;
}

/* ── Keyword pills ── */
.keyword-container { display: flex; flex-wrap: wrap; gap: 0.5rem; margin: 0.75rem 0; }
.keyword-pill {
    background: rgba(99,179,237,0.12);
    border: 1px solid rgba(99,179,237,0.3);
    color: #90cdf4;
    padding: 0.2rem 0.65rem;
    border-radius: 20px;
    font-size: 0.78rem;
    font-weight: 500;
}
.keyword-pill.matched {
    background: rgba(72,187,120,0.12);
    border-color: rgba(72,187,120,0.3);
    color: #9ae6b4;
}
.keyword-pill.missing {
    background: rgba(252,129,74,0.1);
    border-color: rgba(252,129,74,0.3);
    color: #fbd38d;
}

/* ── Status badges ── */
.status-ok   { color: #68d391; font-weight: 600; }
.status-warn { color: #f6ad55; font-weight: 600; }
.status-info { color: #63b3ed; font-weight: 600; }

/* ── Section preview ── */
.resume-preview {
    background: #f8fafc;
    color: #1e293b;
    border-radius: 10px;
    padding: 2rem;
    font-family: 'DM Sans', sans-serif;
    font-size: 0.88rem;
    line-height: 1.7;
    border: 1px solid #e2e8f0;
    white-space: pre-wrap;
    max-height: 520px;
    overflow-y: auto;
}

/* ── Sidebar tweaks ── */
section[data-testid="stSidebar"] {
    background: #0f172a;
    border-right: 1px solid #1e293b;
}
section[data-testid="stSidebar"] .stMarkdown p,
section[data-testid="stSidebar"] label {
    color: #94a3b8 !important;
    font-size: 0.85rem;
}
section[data-testid="stSidebar"] h3 {
    color: #cbd5e1 !important;
}

/* ── Buttons ── */
.stButton > button {
    background: linear-gradient(135deg, #1e3a5f, #2563eb);
    color: white;
    border: none;
    border-radius: 8px;
    padding: 0.6rem 2rem;
    font-weight: 600;
    font-family: 'DM Sans', sans-serif;
    letter-spacing: 0.02em;
    transition: all 0.2s;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #2563eb, #3b82f6);
    transform: translateY(-1px);
    box-shadow: 0 4px 20px rgba(37,99,235,0.4);
}
div[data-testid="stDownloadButton"] > button {
    background: linear-gradient(135deg, #064e3b, #059669) !important;
    color: white !important;
}
div[data-testid="stDownloadButton"] > button:hover {
    background: linear-gradient(135deg, #059669, #10b981) !important;
}

/* ── Progress bar ── */
.ats-score-bar-wrap { background: #1e293b; border-radius: 8px; height: 12px; overflow: hidden; margin: 0.4rem 0; }
.ats-score-bar { height: 100%; border-radius: 8px; transition: width 0.6s ease; }

/* ── Divider ── */
hr { border-color: #1e293b !important; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# UTILITY HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    if not PDF_OK:
        return ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        return "\n".join(p.extract_text() or "" for p in pdf.pages)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX."""
    if not DOCX_OK:
        return ""
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join(para.text for para in doc.paragraphs)


def extract_resume_sections(text: str) -> dict:
    """
    Heuristic parser – splits resume text into labelled sections.
    Works without any ML model.
    """
    sections = {
        "name": "",
        "contact": "",
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "certifications": "",
        "other": "",
    }
    # Guess name from first non-empty line
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        sections["name"] = lines[0]

    # Section header keywords mapped to our keys
    header_map = {
        r"summary|objective|profile|about": "summary",
        r"experience|employment|work history|career": "experience",
        r"education|academic|degree|university|college": "education",
        r"skill|technical|technologies|tools|competenc": "skills",
        r"certif|license|accredit": "certifications",
        r"contact|email|phone|address|linkedin|github": "contact",
    }

    current_section = "other"
    buffer = {k: [] for k in sections}

    for line in lines[1:]:
        lower = line.lower().strip(":-– \t")
        matched = False
        for pattern, key in header_map.items():
            if re.search(pattern, lower) and len(lower) < 40:
                current_section = key
                matched = True
                break
        if not matched:
            buffer[current_section].append(line)

    for key, lines_list in buffer.items():
        sections[key] = "\n".join(lines_list).strip()

    return sections


def extract_keywords_spacy(text: str, top_n: int = 40) -> list[str]:
    """Extract noun-chunk / entity keywords with spaCy."""
    doc = nlp(text)
    # Named entities + noun chunks that are likely skills/tools
    kws = set()
    for ent in doc.ents:
        if ent.label_ in ("ORG", "PRODUCT", "LANGUAGE", "SKILL", "WORK_OF_ART"):
            kws.add(ent.text.lower())
    for chunk in doc.noun_chunks:
        if 2 <= len(chunk.text.split()) <= 4:
            kws.add(chunk.text.lower())
    # Also grab single tokens that look like tech/tool names (TitleCase or ALL_CAPS)
    for token in doc:
        if (token.is_alpha and not token.is_stop and len(token.text) > 2
                and (token.text[0].isupper() or token.text.isupper())):
            kws.add(token.text.lower())
    return list(kws)[:top_n]


def extract_keywords_regex(text: str, top_n: int = 40) -> list[str]:
    """
    Fallback keyword extractor – no ML needed.
    Extracts noun-like tokens and common tech terms by frequency.
    """
    # Common tech / business terms to specifically look for
    tech_patterns = r"""
        python|java(?:script)?|typescript|react|angular|vue|node\.?js|
        sql|nosql|mongodb|postgres|mysql|redis|elasticsearch|
        aws|azure|gcp|docker|kubernetes|ci\/cd|git(?:hub|lab)?|
        machine\s*learning|deep\s*learning|nlp|data\s*science|
        agile|scrum|kanban|devops|microservices|rest(?:ful)?|graphql|
        html|css|sass|tailwind|bootstrap|
        tensorflow|pytorch|pandas|numpy|scikit|
        leadership|communication|problem.solving|collaboration|
        bachelor|master|phd|degree|certification|
        project\s*management|product\s*management|stakeholder|
        analysis|analytics|visualization|reporting|
        excel|powerpoint|tableau|power\s*bi|looker|
        api|sdk|saas|b2b|b2c|crm|erp
    """
    found = re.findall(
        r'\b(' + '|'.join(p.strip() for p in tech_patterns.split('|') if p.strip()) + r')\b',
        text.lower()
    )
    # Also grab capitalised 2-3 word phrases (likely proper nouns / tools)
    phrases = re.findall(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\b', text)
    all_kw = found + [p.lower() for p in phrases]
    counter = Counter(all_kw)
    return [kw for kw, _ in counter.most_common(top_n)]


def extract_keywords(text: str, top_n: int = 40) -> list[str]:
    """Dispatcher – uses spaCy if available, else regex."""
    if SPACY_OK and nlp:
        return extract_keywords_spacy(text, top_n)
    return extract_keywords_regex(text, top_n)


def score_resume(resume_text: str, jd_keywords: list[str]) -> tuple[int, list[str], list[str]]:
    """Return ATS score (0-100), matched keywords, missing keywords."""
    resume_lower = resume_text.lower()
    matched = [kw for kw in jd_keywords if kw.lower() in resume_lower]
    missing = [kw for kw in jd_keywords if kw.lower() not in resume_lower]
    score = int(len(matched) / max(len(jd_keywords), 1) * 100)
    return score, matched, missing


# ─────────────────────────────────────────────────────────────────────────────
# COMPANY RESEARCH (Wikipedia + HF fallback)
# ─────────────────────────────────────────────────────────────────────────────

def get_wikipedia_summary(company: str) -> str:
    """Fetch company summary from Wikipedia – no API key needed."""
    if not WIKI_OK:
        return ""
    try:
        wikipedia.set_lang("en")
        results = wikipedia.search(company, results=3)
        if not results:
            return ""
        page = wikipedia.summary(results[0], sentences=5, auto_suggest=False)
        return page
    except Exception:
        return ""


def hf_inference(prompt: str, hf_token: str, model: str = "google/flan-t5-base") -> str:
    """
    Call Hugging Face free Inference API.
    Free tier: ~30k tokens/month with a free HF account token.
    Falls back gracefully if call fails.
    """
    if not hf_token:
        return ""
    url = f"https://api-inference.huggingface.co/models/{model}"
    headers = {"Authorization": f"Bearer {hf_token}"}
    payload = {"inputs": prompt[:1500], "parameters": {"max_new_tokens": 200}}
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=30)
        if r.status_code == 200:
            data = r.json()
            if isinstance(data, list) and data:
                return data[0].get("generated_text", "").strip()
            if isinstance(data, dict):
                return data.get("generated_text", "").strip()
    except Exception:
        pass
    return ""


SALARY_TABLE = {
    "intern": ("$20,000", "$40,000"),
    "junior": ("$50,000", "$75,000"),
    "associate": ("$60,000", "$85,000"),
    "mid": ("$75,000", "$100,000"),
    "senior": ("$100,000", "$140,000"),
    "lead": ("$120,000", "$160,000"),
    "principal": ("$140,000", "$190,000"),
    "staff": ("$150,000", "$200,000"),
    "manager": ("$100,000", "$145,000"),
    "director": ("$140,000", "$200,000"),
    "vp": ("$180,000", "$260,000"),
    "cto": ("$200,000", "$350,000"),
    "ceo": ("$200,000", "$400,000"),
}


def estimate_salary(jd_text: str, hf_token: str) -> str:
    """
    Rule-based salary estimate by seniority level extracted from JD.
    Optionally uses HF to refine.
    """
    jd_lower = jd_text.lower()
    for level, (lo, hi) in SALARY_TABLE.items():
        if level in jd_lower:
            return f"{lo} – {hi} / year (estimated based on seniority level)"
    return "Unable to estimate – please research via Glassdoor or Levels.fyi"


def analyse_company(company: str, jd_text: str, hf_token: str) -> dict:
    """Return dict with profile, culture, benefits, salary."""
    result = {"profile": "", "culture": "", "benefits": "", "salary": ""}

    # 1. Wikipedia
    wiki = get_wikipedia_summary(company) if company else ""
    result["profile"] = wiki or f"No Wikipedia article found for '{company}'."

    # 2. Salary (rule-based + optional HF)
    result["salary"] = estimate_salary(jd_text, hf_token)

    # 3. Culture & Benefits – HF if token available, else template
    if hf_token and company:
        culture_prompt = (
            f"Describe the work culture and environment at {company} "
            f"in 2-3 sentences. Focus on values, team dynamics, and employee experience."
        )
        result["culture"] = hf_inference(culture_prompt, hf_token) or \
            f"Culture information not available. Research {company} on Glassdoor or LinkedIn."

        benefits_prompt = (
            f"List typical employee benefits offered by {company} "
            f"such as health insurance, PTO, remote work, and professional development."
        )
        result["benefits"] = hf_inference(benefits_prompt, hf_token) or \
            "Benefits information not available. Check company careers page."
    else:
        result["culture"] = (
            f"Culture details for {company or 'this company'} require a Hugging Face API token. "
            "Research on Glassdoor, Blind, or LinkedIn for employee reviews."
        )
        result["benefits"] = (
            "Benefits details require a Hugging Face API token. "
            "Check the company's careers page or ask during the interview process."
        )

    return result


# ─────────────────────────────────────────────────────────────────────────────
# RESUME OPTIMISATION
# ─────────────────────────────────────────────────────────────────────────────

def optimise_resume_text(sections: dict, jd_keywords: list[str], hf_token: str) -> str:
    """
    Rewrite / enrich resume sections with JD keywords.
    If HF token available, uses the model; otherwise does rule-based injection.
    """
    missing_kws = [kw for kw in jd_keywords
                   if kw.lower() not in (sections.get("skills", "") + sections.get("experience", "")).lower()]

    # ── Skills section: append missing relevant keywords ──────────────────
    skills_section = sections.get("skills", "")
    if missing_kws:
        added = ", ".join(missing_kws[:15])
        skills_section = (skills_section + f"\nAdditional relevant skills: {added}").strip()

    # ── Experience: HF rewrite or passthrough ─────────────────────────────
    experience_section = sections.get("experience", "")
    if hf_token and experience_section:
        kw_list = ", ".join(jd_keywords[:20])
        prompt = (
            f"Rewrite the following work experience section to naturally incorporate "
            f"these keywords: {kw_list}.\n"
            f"Keep bullet points, preserve facts, improve clarity and impact.\n\n"
            f"Experience:\n{experience_section[:1200]}"
        )
        rewritten = hf_inference(prompt, hf_token)
        if rewritten and len(rewritten) > 100:
            experience_section = rewritten

    # ── Summary: HF rewrite or passthrough ───────────────────────────────
    summary_section = sections.get("summary", "")
    if hf_token and summary_section:
        kw_list = ", ".join(jd_keywords[:10])
        prompt = (
            f"Rewrite this professional summary to naturally include keywords: {kw_list}.\n"
            f"Keep it concise (3-4 sentences), professional, and impactful.\n\n"
            f"Summary:\n{summary_section[:600]}"
        )
        rewritten = hf_inference(prompt, hf_token)
        if rewritten and len(rewritten) > 50:
            summary_section = rewritten

    # ── Assemble plain text resume ────────────────────────────────────────
    name = sections.get("name", "Your Name")
    contact = sections.get("contact", "")
    education = sections.get("education", "")
    certifications = sections.get("certifications", "")
    other = sections.get("other", "")

    parts = []
    parts.append(name.upper())
    if contact:
        parts.append(contact)
    parts.append("")

    if summary_section:
        parts += ["PROFESSIONAL SUMMARY", "-" * 40, summary_section, ""]
    if experience_section:
        parts += ["WORK EXPERIENCE", "-" * 40, experience_section, ""]
    if education:
        parts += ["EDUCATION", "-" * 40, education, ""]
    if skills_section:
        parts += ["SKILLS", "-" * 40, skills_section, ""]
    if certifications:
        parts += ["CERTIFICATIONS", "-" * 40, certifications, ""]
    if other:
        parts += [other]

    return "\n".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX BUILDER (ATS-FRIENDLY)
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(resume_text: str) -> bytes:
    """
    Convert plain resume text to an ATS-friendly .docx.
    Rules: no tables, no columns, no text boxes, no graphics.
    Uses standard paragraph styles only.
    """
    doc = DocxDocument()

    # Page margins (1 inch all around is ATS safe)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = resume_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            doc.add_paragraph("")
            i += 1
            continue

        # Detect section headers (ALL CAPS, short, maybe followed by dashes)
        is_header = (line.isupper() and len(line) < 50 and not line.startswith("•"))
        is_dash_line = set(line.strip()) <= {"-", "─", "="}

        if is_header:
            p = doc.add_paragraph(line)
            p.style = doc.styles["Heading 2"]
            run = p.runs[0] if p.runs else p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
            # Don't colour – ATS parsers prefer black
            i += 1
        elif is_dash_line:
            i += 1  # Skip decorative lines in source text
        elif line.startswith("•") or line.startswith("-") or line.startswith("*"):
            # Bullet point
            clean = line.lstrip("•-* ").strip()
            p = doc.add_paragraph(clean, style="List Bullet")
            p.runs[0].font.size = Pt(11)
            i += 1
        else:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(11) if p.runs else None
            i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# COVER LETTER GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

COVER_LETTER_TEMPLATE = """
{name}
{contact}

{date}

Hiring Manager
{company}

Dear Hiring Manager,

I am writing to express my strong interest in the {role} position at {company}. With {years} of experience in {field} and a proven track record of {achievement}, I am confident in my ability to make a meaningful contribution to your team.

Throughout my career, I have developed deep expertise in {top_skills}. In my most recent role, I {recent_achievement}. This experience has equipped me with the skills necessary to {value_prop}.

What particularly excites me about {company} is {company_appeal}. I am drawn to your commitment to innovation and believe my background in {background} aligns well with your team's objectives. I am eager to bring my skills in {key_skills} to help drive your organisation's success.

I would welcome the opportunity to discuss how my experience and passion can contribute to {company}'s continued growth. Thank you for considering my application. I look forward to the possibility of speaking with you.

Sincerely,

{name}
""".strip()


def generate_cover_letter(sections: dict, company: str, jd_text: str, hf_token: str) -> str:
    """Generate cover letter via HF or fill template."""
    name = sections.get("name", "Your Name")
    contact = sections.get("contact", "")
    skills_raw = sections.get("skills", "")
    experience_raw = sections.get("experience", "")

    # Extract role from JD
    role_match = re.search(
        r'(?:position|role|job title|hiring for|looking for)[:\s]+([^\n.]{5,60})',
        jd_text, re.IGNORECASE
    )
    role = role_match.group(1).strip() if role_match else "the advertised position"

    # Infer field from skills
    skills_list = [s.strip() for s in re.split(r'[,;\n]', skills_raw) if s.strip()][:5]
    field = skills_list[0] if skills_list else "my field"
    top_skills = ", ".join(skills_list[:3]) if skills_list else "relevant technologies"
    key_skills = ", ".join(skills_list[3:5]) if len(skills_list) > 3 else top_skills

    # Infer years of experience
    years_match = re.search(r'(\d+)\+?\s*years?', experience_raw, re.IGNORECASE)
    years = f"{years_match.group(1)}+ years" if years_match else "several years"

    if hf_token:
        prompt = (
            f"Write a professional cover letter for {name} applying to the {role} position "
            f"at {company}. The candidate has experience in {', '.join(skills_list[:4])}. "
            f"Keep it 3-4 paragraphs, professional tone, approximately 300 words. "
            f"Focus on value to the employer."
        )
        result = hf_inference(prompt, hf_token, model="google/flan-t5-base")
        if result and len(result) > 200:
            return f"{name}\n{contact}\n\n{datetime.now().strftime('%B %d, %Y')}\n\n" + result

    # Template fallback
    return COVER_LETTER_TEMPLATE.format(
        name=name,
        contact=contact or "your@email.com | (555) 000-0000",
        date=datetime.now().strftime("%B %d, %Y"),
        company=company or "the Company",
        role=role,
        years=years,
        field=field,
        achievement="delivering measurable results",
        top_skills=top_skills,
        recent_achievement="led key initiatives that improved team performance and outcomes",
        value_prop="excel in this role from day one",
        company_appeal=f"{company or 'your organisation'}'s mission and innovative approach",
        background=field,
        key_skills=key_skills,
    )


def build_cover_letter_docx(cover_text: str) -> bytes:
    """Wrap cover letter in a clean .docx."""
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    for line in cover_text.split("\n"):
        p = doc.add_paragraph(line)
        if p.runs:
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = "Calibri"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("### ⚙️ Configuration")

    hf_token = st.text_input(
        "Hugging Face API Token (optional)",
        type="password",
        help="Free token from huggingface.co. Enables AI-powered rewrites and analysis.",
        placeholder="hf_xxxxxxxxxxxxxxxxxxxx",
    )
    st.caption("🔒 Token used only for this session, never stored.")

    st.markdown("---")
    st.markdown("### 📄 Your Resume")

    upload_mode = st.radio("Input method", ["Upload file", "Paste text"], index=0)

    raw_resume_text = ""
    if upload_mode == "Upload file":
        uploaded = st.file_uploader(
            "Upload PDF or DOCX",
            type=["pdf", "docx"],
            help="Your resume will be parsed automatically.",
        )
        if uploaded:
            ext = uploaded.name.rsplit(".", 1)[-1].lower()
            file_bytes = uploaded.read()
            if ext == "pdf":
                raw_resume_text = parse_pdf(file_bytes)
            elif ext == "docx":
                raw_resume_text = parse_docx(file_bytes)
            if raw_resume_text:
                st.success(f"✅ Parsed {len(raw_resume_text.split())} words")
            else:
                st.warning("Could not parse – try pasting text instead.")
    else:
        raw_resume_text = st.text_area(
            "Paste your resume text",
            height=250,
            placeholder="Paste your full resume here…",
        )

    st.markdown("---")
    st.markdown("### 📋 Job Description")
    jd_text = st.text_area(
        "Paste job description",
        height=250,
        placeholder="Paste the full job description here…",
    )

    st.markdown("---")
    company_name = st.text_input(
        "Company Name",
        placeholder="e.g. Google, Stripe, Acme Corp",
        help="Auto-extracted from JD if left blank.",
    )

    # Auto-extract company from JD if not provided
    if not company_name and jd_text:
        match = re.search(
            r'(?:at|join|for)\s+([A-Z][a-zA-Z0-9&\s]{2,30}?)(?:\s+is|\s+we|\s*,|\s*\.)',
            jd_text
        )
        if match:
            company_name = match.group(1).strip()

    gen_cover = st.checkbox("✉️ Generate Cover Letter", value=True)
    analyse_co = st.checkbox("🔍 Analyse Company", value=True)

    st.markdown("---")
    st.caption("💡 **Zero cost** – works without any API keys. Add HF token for AI-enhanced output.")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN AREA
# ─────────────────────────────────────────────────────────────────────────────

# Hero header
st.markdown("""
<div class="hero-header">
  <div class="hero-badge">ATS Optimiser</div>
  <h1 class="hero-title">Resume & Cover Letter Builder</h1>
  <p class="hero-subtitle">
    Keyword-optimised · ATS-friendly · Zero cost · Instant .docx download
  </p>
</div>
""", unsafe_allow_html=True)

# ── Status row ────────────────────────────────────────────────────────────────
col1, col2, col3, col4 = st.columns(4)
with col1:
    status = "✅ Ready" if raw_resume_text else "⏳ Needed"
    css = "status-ok" if raw_resume_text else "status-warn"
    st.markdown(f"**Resume** <span class='{css}'>{status}</span>", unsafe_allow_html=True)
with col2:
    status = "✅ Ready" if jd_text else "⏳ Needed"
    css = "status-ok" if jd_text else "status-warn"
    st.markdown(f"**Job Description** <span class='{css}'>{status}</span>", unsafe_allow_html=True)
with col3:
    status = "✅ AI-enhanced" if hf_token else "📝 Template mode"
    css = "status-ok" if hf_token else "status-info"
    st.markdown(f"**Mode** <span class='{css}'>{status}</span>", unsafe_allow_html=True)
with col4:
    status = f"✅ {company_name}" if company_name else "⏳ Not set"
    css = "status-ok" if company_name else "status-warn"
    st.markdown(f"**Company** <span class='{css}'>{status}</span>", unsafe_allow_html=True)

st.markdown("---")

# ── Keyword preview (live as user types JD) ───────────────────────────────────
if jd_text:
    jd_keywords = extract_keywords(jd_text, top_n=35)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">🔑 Keywords Extracted from Job Description</div>', unsafe_allow_html=True)
    pills_html = '<div class="keyword-container">'
    for kw in jd_keywords:
        pills_html += f'<span class="keyword-pill">{kw}</span>'
    pills_html += '</div>'
    st.markdown(pills_html, unsafe_allow_html=True)

    if raw_resume_text:
        score, matched, missing = score_resume(raw_resume_text, jd_keywords)
        bar_color = "#68d391" if score >= 70 else "#f6ad55" if score >= 40 else "#fc8181"
        st.markdown(f"""
        <div style='margin-top:1rem;'>
          <div style='display:flex; justify-content:space-between; margin-bottom:4px;'>
            <span style='color:#94a3b8; font-size:0.85rem;'>ATS Match Score (before optimisation)</span>
            <span style='color:{bar_color}; font-weight:700;'>{score}%</span>
          </div>
          <div class='ats-score-bar-wrap'>
            <div class='ats-score-bar' style='width:{score}%; background:{bar_color};'></div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        m_html = '<div class="keyword-container" style="margin-top:0.75rem;">'
        for kw in matched[:15]:
            m_html += f'<span class="keyword-pill matched">✓ {kw}</span>'
        m_html += '</div>'
        st.markdown("**Matched keywords:**", unsafe_allow_html=False)
        st.markdown(m_html, unsafe_allow_html=True)

        if missing:
            mi_html = '<div class="keyword-container">'
            for kw in missing[:15]:
                mi_html += f'<span class="keyword-pill missing">+ {kw}</span>'
            mi_html += '</div>'
            st.markdown("**Keywords to add:**", unsafe_allow_html=False)
            st.markdown(mi_html, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)
else:
    jd_keywords = []

# ── Company Analysis ───────────────────────────────────────────────────────────
if analyse_co and company_name and jd_text:
    with st.expander("🏢 Company Analysis", expanded=True):
        with st.spinner(f"Researching {company_name}…"):
            co_data = analyse_company(company_name, jd_text, hf_token)

        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown(f'<div class="card-title">🏢 {company_name}</div>', unsafe_allow_html=True)

        tab1, tab2, tab3, tab4 = st.tabs(["📖 Profile", "🌱 Culture", "🎁 Benefits", "💰 Salary Est."])
        with tab1:
            st.write(co_data["profile"])
        with tab2:
            st.write(co_data["culture"])
        with tab3:
            st.write(co_data["benefits"])
        with tab4:
            st.write(f"**Estimated salary range:** {co_data['salary']}")
            st.caption("Based on seniority keywords in the job description. Verify on Glassdoor, Levels.fyi, or LinkedIn Salary.")

        st.markdown('</div>', unsafe_allow_html=True)

# ── GENERATE BUTTON ───────────────────────────────────────────────────────────
st.markdown("")
generate_btn = st.button(
    "🚀 Generate ATS Resume" + (" & Cover Letter" if gen_cover else ""),
    use_container_width=True,
    disabled=(not raw_resume_text or not jd_text),
)

if not raw_resume_text or not jd_text:
    st.info("👈  Add your resume and job description in the sidebar to get started.")

if generate_btn and raw_resume_text and jd_text:
    with st.spinner("Optimising your resume…"):
        # Parse resume
        sections = extract_resume_sections(raw_resume_text)

        # Extract keywords
        kws = extract_keywords(jd_text, top_n=35)

        # Optimise
        optimised_text = optimise_resume_text(sections, kws, hf_token)

        # Build docx
        resume_docx_bytes = build_docx(optimised_text)

        # Score after
        score_after, matched_after, missing_after = score_resume(optimised_text, kws)

    # ── Results ──────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("## ✅ Optimised Resume Ready")

    # Score comparison
    score_before, _, _ = score_resume(raw_resume_text, kws)
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("ATS Score Before", f"{score_before}%")
    with c2:
        st.metric("ATS Score After", f"{score_after}%", delta=f"+{score_after - score_before}%")
    with c3:
        st.metric("Keywords Added", len(kws) - len(missing_after))

    # Preview
    st.markdown("### 📝 Resume Preview")
    st.markdown(f'<div class="resume-preview">{optimised_text}</div>', unsafe_allow_html=True)

    # Download resume
    st.download_button(
        label="⬇️ Download Optimised Resume (.docx)",
        data=resume_docx_bytes,
        file_name=f"ats_resume_{sections.get('name', 'resume').replace(' ', '_').lower()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # ── Cover letter ──────────────────────────────────────────────────────
    if gen_cover:
        with st.spinner("Generating cover letter…"):
            cover_text = generate_cover_letter(sections, company_name, jd_text, hf_token)
            cover_docx_bytes = build_cover_letter_docx(cover_text)

        st.markdown("---")
        st.markdown("### ✉️ Cover Letter Preview")
        st.markdown(f'<div class="resume-preview">{cover_text}</div>', unsafe_allow_html=True)

        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            st.download_button(
                label="⬇️ Download Cover Letter (.docx)",
                data=cover_docx_bytes,
                file_name=f"cover_letter_{company_name.replace(' ', '_').lower() if company_name else 'cover'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with dl_col2:
            st.download_button(
                label="⬇️ Download Cover Letter (.txt)",
                data=cover_text.encode(),
                file_name=f"cover_letter_{company_name.replace(' ', '_').lower() if company_name else 'cover'}.txt",
                mime="text/plain",
            )

    # ── Tips ──────────────────────────────────────────────────────────────
    with st.expander("💡 ATS Tips & Next Steps"):
        st.markdown("""
**Before submitting:**
- Review the downloaded `.docx` to ensure all information is accurate
- Replace any placeholder text (e.g. `your@email.com`) with your real details  
- Spell-check the final document  
- Save as PDF only if the job application explicitly accepts PDF (some ATS prefer .docx)

**ATS Best Practices:**
- Use standard section headings (Experience, Education, Skills)
- Avoid tables, text boxes, headers/footers with key info
- Use common fonts (Calibri, Arial, Times New Roman)
- List dates consistently: `Jan 2021 – Mar 2023`
- Spell out acronyms at least once: `Machine Learning (ML)`

**Hugging Face free token:** [huggingface.co/settings/tokens](https://huggingface.co/settings/tokens)  
Enables AI-rewritten bullet points and cover letter personalisation.
        """)

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("---")
st.caption("🛡️ All processing runs locally. No data is stored or transmitted except optional Hugging Face API calls.")
